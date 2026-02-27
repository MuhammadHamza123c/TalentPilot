import requests
import pdfplumber
from docx import Document
from io import BytesIO

# PDF extractor from file-like object
def extract_pdf_from_url(url: str) -> str:
    text = ""
    try:
        response = requests.get(url)
        response.raise_for_status()  # ensure we got the file
        with pdfplumber.open(BytesIO(response.content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if not text.strip():
            text = 'Upload Correct Resume'
    except Exception as e:
        raise ValueError(f"Failed to extract PDF text: {str(e)}")
    return text.strip()


# DOCX extractor from file-like object
def extract_docx_from_url(url: str) -> str:
    text = []
    try:
        response = requests.get(url)
        response.raise_for_status()
        doc = Document(BytesIO(response.content))

        # Extract paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                text.append(p.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text.append(row_text)

        # Extract headers and footers
        for section in doc.sections:
            header = section.header
            for p in header.paragraphs:
                if p.text.strip():
                    text.append(p.text)
            for table in header.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text.append(row_text)

            footer = section.footer
            for p in footer.paragraphs:
                if p.text.strip():
                    text.append(p.text)
            for table in footer.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text.append(row_text)

        if not text:
            text.append('Upload Correct Resume')

    except Exception as e:
        raise ValueError(f"Failed to extract DOCX text: {str(e)}")

    return "\n".join(text).strip()


